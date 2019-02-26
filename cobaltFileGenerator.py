import os
import sys
from docx import Document

output_dir = "generated-test-files"
MAGIC_NUMBERS = {
    "jpg": bytearray.fromhex("FF D8 FF DB"),
    "jpeg": bytearray.fromhex("FF D8 FF EE"),
    "pdf": bytearray.fromhex("25 50 44 46 2d"),
    "png": bytearray.fromhex("89 50 4E 47 0D 0A 1A 0A"),
    "docx": bytearray.fromhex(
        "50 4B 03 04 14 00 00 00 08 00 6F 6E 74 65 6E 74 5F 54 79"),
    "doc": bytearray.fromhex("D0 CF 11 E0 A1 B1 1A E1"),
    "gif": bytearray.fromhex("47 49 46 38 37 61")
}

def _create_doc(file_name):
    document = Document()
    document.add_heading("Test Heading", 0)
    document.add_paragraph("Test Paragraph")
    document.save(file_name)


def create_file(file_ext, out_dir):
    if ext not in MAGIC_NUMBERS.keys():
        raise ValueError("Unsupported extension {}".format(file_ext))

    file_name = "{0}/{1}TestFile.{2}".format(out_dir, file_ext, file_ext)

    if ext in ("docx", "doc"):
        _create_doc(file_name)
        return

    with open(file_name, "wb") as fh:
        fh.write(MAGIC_NUMBERS[ext])


def ensure_output_dir(out_dir):
    if os.path.exists(out_dir):
        os.system("rm -rf {}".format(out_dir))

    os.mkdir(out_dir)


if __name__ == "__main__":
    if len(sys.argv) > 1:
        extensions = sys.argv[1:]
    else:
        extensions = MAGIC_NUMBERS.keys()

    ensure_output_dir(output_dir)

    for ext in extensions:
        create_file(ext, output_dir)
